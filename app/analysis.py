from flask import Blueprint, render_template, request, url_for, session, after_this_request, send_file, redirect
import google.generativeai as genai
from PIL import Image
import logging
import cv2
import io
import os
import requests
import whisper
from moviepy.editor import VideoFileClip
from concurrent.futures import ThreadPoolExecutor, as_completed
from werkzeug.utils import secure_filename
import time
import google.api_core.exceptions
from docx import Document
from dotenv import load_dotenv
from .gemini_api_manager import api_manager
from multiprocessing import Pool

# Load environment variables
load_dotenv()

# Create a Blueprint for the analysis routes
analysis_bp = Blueprint('analysis_bp', __name__)

# Load Whisper model for audio transcription
whisper_model = whisper.load_model("base")

@analysis_bp.route('/')
def index():
    if 'username' in session:
        return render_template('index.html')
    return redirect(url_for('auth_bp.login'))

@analysis_bp.route("/analyze", methods=["POST"])
def analyze():
    logging.info("🔍 Received /analyze request")
    file_type = request.form.get("file_type")

    if file_type == "image":
        return analyze_image()
    elif file_type == "audio":
        return analyze_audio()
    elif file_type == "video":
        video_file = request.files.get("video")
        return analyze_video(video_file)

    return render_template("result.html", summary="❌ Invalid file type."), 400

def analyze_image():
    file = request.files.get("image")
    image_url = request.form.get("image_url")
    image_bytes = None
    image_display_path = None

    # --- File Upload/URL Handling (this part remains largely the same) ---
    if file and file.filename != "":
        filename = secure_filename(file.filename)
        file_path = os.path.join(analysis_bp.root_path, '..', 'static', 'uploads', filename)
        file.save(file_path)
        with open(file_path, "rb") as f:
            image_bytes = f.read()
        image_display_path = url_for('static', filename=f'uploads/{filename}')

    elif image_url and image_url.strip() != "":
        try:
            response = requests.get(image_url, timeout=10)
            response.raise_for_status()
            content_type = response.headers.get("Content-Type", "")
            if not content_type.startswith("image/"):
                return render_template("result.html", summary="❌ URL does not point to an image."), 400
            image_bytes = response.content
            filename = os.path.basename(image_url.split("?")[0])
            if not any(filename.lower().endswith(ext) for ext in [".jpg", ".jpeg", ".png", ".gif", ".webp"]):
                filename += ".jpg"
            filename = secure_filename(filename)
            file_path = os.path.join(analysis_bp.root_path, '..', 'static', 'uploads', filename)
            with open(file_path, "wb") as f:
                f.write(image_bytes)
            image_display_path = url_for('static', filename=f'uploads/{filename}')
        except Exception as e:
            return render_template("result.html", summary=f"❌ Failed to fetch image: {str(e)}"), 400
    else:
        return render_template("result.html", summary="❌ No image or URL provided."), 400

    try:
        image = Image.open(io.BytesIO(image_bytes))
    except Exception as e:
        return render_template("result.html", summary=f"❌ Unable to process image: {str(e)}"), 400

    # --- Part 1: Classify the Image with Retries ---
    type_prompt = """
    You are a content classification expert. Given an image, classify it into **one** of the following exact categories only:

    People, Scenery, Graphs, UI, Text, ID, Product, Animals, Objects, Others

    Instructions:
    - Return just **one** of the above words.
    - Do NOT use any symbols, markdown, or explanation.
    - If the image doesn’t clearly fall into any category, return: Others
    """
    imgtype = ""
    for attempt in range(5):
        try:
            genai.configure(api_key=api_manager.current_key)
            model = genai.GenerativeModel("models/gemini-2.0-flash")
            imgtype_response = model.generate_content([image, type_prompt])
            imgtype = imgtype_response.text.strip().lower()
            break  # Break the loop on success
        except google.api_core.exceptions.ResourceExhausted:
            logging.error("⚠️ Rate limit hit during image classification. Switching key and retrying...")
            api_manager.get_next_key()
            time.sleep(1)
        except Exception as e:
            logging.error(f"❌ Unexpected error during image classification: {e}")
            return render_template("result.html", summary=f"❌ An error occurred: {str(e)}"), 500
    else:
        return render_template("result.html", summary="❌ Failed to classify image after multiple attempts due to API issues."), 500

    # --- Part 2: Generate the Summary with Retries based on classification ---
    custom_prompt = """
    You are an expert visual analyst. Describe the image as clearly and objectively as possible. Focus on:
    - Key objects or subjects
    - Any notable visual features
    - Try to infer the likely purpose, context, or mood of the image
    Avoid formatting (no markdown, no symbols). Keep it clean and structured.
    """

    # (The rest of your conditional prompts for image analysis)
    if imgtype == "people":
        custom_prompt += """
        This image shows people or portraits. Describe:
        - Number of individuals, gender, and age estimates
        - Clothing and facial expressions
        - Activity, posture, and background setting
        - Possible social setting or context
        """
    elif imgtype == "scenery":
        custom_prompt += """
        This image shows scenery or a landscape. Describe:
        - Natural or urban environment
        - Weather, lighting, time of day
        - Key objects (trees, buildings, roads, etc.)
        - Emotional or aesthetic tone
        """
    elif imgtype == "graphs":
        custom_prompt += """
        This image contains a graph or chart. Describe:
        - Type of graph (bar, line, pie, etc.)
        - Data trends or key observations
        - Any visible axes or labels
        - Likely domain (finance, health, education, etc.)
        """
    elif imgtype == "ui":
        custom_prompt += """
        This is a user interface screenshot. Describe:
        - Type of UI (website, app, dashboard)
        - Visible UI elements (buttons, tables, inputs, charts)
        - Branding, color scheme, or theme
        - Inferred purpose or user goal
        """
    elif imgtype == "text":
        custom_prompt += """
        This image shows text content. Describe:
        - Main headings or paragraph structure
        - Language and tone
        - Formatting (font size, bold/italic, layout)
        - What the text is trying to communicate
        """
    elif imgtype == "id":
        custom_prompt += """
        This image shows an ID or document. Describe:
        - Type of document (passport, Aadhar, license, etc.)
        - Layout: photo placement, fields (name, ID, DOB)
        - Any security features, logos, or text
        - Purpose of the document
        """
    elif imgtype == "product":
        custom_prompt += """
        This image shows a product. Describe:
        - Product type and appearance
        - Brand logos or labels
        - Background setup (studio, plain, natural)
        - Inferred use-case or target customer
        """
    elif imgtype == "animals":
        custom_prompt += """
        This image shows animals. Describe:
        - Type of animal(s) (dog, cat, bird, etc.)
        - Number, appearance, and actions
        - Environment (indoor, outdoor, zoo, wild)
        - Mood or tone of the image
        """
    elif imgtype == "objects":
        custom_prompt += """
        This image shows physical objects. Describe:
        - Type of object(s) (electronics, tools, toys, etc.)
        - Materials, shape, and condition
        - Background or context clues
        - Possible use or function of the objects
        """
    elif imgtype == "others":
        custom_prompt += """
        Please:
        - Describe visible elements in detail
        - Mention key shapes, colors, patterns, and textures
        - Infer possible purpose or meaning
        - Keep it readable and structured
        """

    response = None
    for attempt in range(5):
        try:
            genai.configure(api_key=api_manager.current_key)
            model = genai.GenerativeModel("models/gemini-2.0-flash")
            response = model.generate_content([image, custom_prompt])
            break  # Break the loop on success
        except google.api_core.exceptions.ResourceExhausted:
            logging.error("⚠️ Rate limit hit during image summary. Switching key and retrying...")
            api_manager.get_next_key()
            time.sleep(1)
        except Exception as e:
            logging.error(f"❌ Unexpected error during image summary: {e}")
            return render_template("result.html", summary=f"❌ An error occurred: {str(e)}"), 500
    else:
        return render_template("result.html", summary="❌ Failed to summarize image after multiple attempts due to API issues."), 500

    session["summary"] = response.text
    return render_template("result.html", summary=response.text, image_url=image_display_path, mode="image")

def analyze_audio():
    audio = request.files.get("audio")
    if not audio or audio.filename == "":
        return render_template("result.html", summary="❌ No audio file uploaded."), 400

    filename = secure_filename(audio.filename)
    audio_path = os.path.join(analysis_bp.root_path, '..', 'static', 'uploads', filename)
    audio.save(audio_path)

    if not os.path.exists(audio_path):
        return render_template("result.html", summary="❌ Audio file was not saved correctly."), 500

    # --- Part 1: Audio Transcription with Whisper (no changes needed here) ---
    try:
        logging.info("🧠 Transcribing audio with Whisper...")
        result = whisper_model.transcribe(audio_path)
        transcription = result["text"]
        logging.info(f"📜 Transcribed Text: {transcription}")
    except Exception as e:
        return render_template("result.html", summary=f"❌ Audio transcription failed: {str(e)}"), 500

    # --- Part 2: Summarization with Gemini API (this part is refactored) ---
    prompt = f"""
    The following is a transcription from an audio file. It may contain profanity or be in a language other than English.

    1. First, **remove any profane or offensive words** by replacing them with [redacted].
    2. Then, **translate the cleaned transcription into English**, if it's in another language.
    3. After cleaning and translating, analyze the content and do the following:

    a. Identify the type of audio:
        - casual conversation
        - formal discussion
        - song or music lyrics
        - speech or monologue
        - other (specify)

    b. Summarize the content in a few clean, structured lines:
        - If it’s a conversation, briefly explain what each person is talking about.
        - If it’s a song, describe the mood and the message.
        - If it’s noisy or unclear, mention that.

    4. Do not list or format dialogues. Just describe the intent or message behind them.
    5. Ensure the summary is clean, neutral, and free from offensive language.
    6. Keep everything concise, readable, and suitable for a general audience.
    7. Output only the final summary — no markdown, no labels, no symbols.

    Transcript:
    {transcription}
    """

    response = None
    for attempt in range(5):
        try:
            # Re-configure the API with the current key from the manager
            genai.configure(api_key=api_manager.current_key)
            model = genai.GenerativeModel("models/gemini-2.0-flash")
            response = model.generate_content(prompt)
            break  # Break the loop on success
        except google.api_core.exceptions.ResourceExhausted:
            logging.error("⚠️ Rate limit hit during audio summarization. Switching key and retrying...")
            api_manager.get_next_key()
            time.sleep(1)  # Small delay before retrying
        except Exception as e:
            logging.error(f"❌ Unexpected error during audio summarization: {e}")
            return render_template("result.html", summary=f"❌ An error occurred during summarization: {str(e)}"), 500
    else:
        # This block runs if the loop finishes without a 'break'
        return render_template("result.html", summary="❌ Failed to summarize audio after multiple attempts due to API issues."), 500

    audio_url = url_for('static', filename=f'uploads/{filename}')
    session["summary"] = response.text
    return render_template("result.html", summary=response.text, audio_url=audio_url, mode="audio")


# --- Helper: Extract frames from video ---
def extract_frames(video_path, interval=90):
    frames = []
    cap = cv2.VideoCapture(video_path)
    frame_count = 0
    success, frame = cap.read()

    while success:
        if frame_count % interval == 0:
            frame_path = os.path.join(os.path.dirname(video_path), f"frame_{frame_count}.jpg")
            cv2.imwrite(frame_path, frame)
            frames.append(frame_path)
        success, frame = cap.read()
        frame_count += 1

    cap.release()
    return frames

# --- Helper: Describe a frame using Gemini ---
def describe_frame(path):
    try:
        image = Image.open(path)
    except Exception as e:
        logging.error(f"❌ Could not open image file: {path} - {e}")
        return f"[Error: Could not open image file: {str(e)}]"

    for attempt in range(5):
        try:
            genai.configure(api_key=api_manager.current_key)
            model = genai.GenerativeModel("models/gemini-2.0-flash")
            prompt = "Describe what's happening in this image."
            response = model.generate_content([prompt, image])
            image.close()
            return response.text
        except google.api_core.exceptions.ResourceExhausted:
            logging.warning("⚠️ Rate limit hit. Switching API key...")
            api_manager.get_next_key()
            time.sleep(1)
        except Exception as e:
            logging.error(f"❌ Unexpected error for image {path}: {e}")
            image.close()
            return f"[Error describing frame: {str(e)}]"

    image.close()
    return "[Failed to describe frame after multiple attempts due to rate limits.]"

# --- Helper: Extract audio and transcribe using Whisper ---
def extract_audio_and_transcribe(video_path):
    try:
        video = VideoFileClip(video_path)
        if video.audio is None:
            logging.warning("⚠️ No audio found in video.")
            video.close()
            return ""

        audio_path = os.path.join(os.path.dirname(video_path), "temp_audio.mp3")
        video.audio.write_audiofile(audio_path, logger=None)
        video.close()

        model = whisper.load_model("base")
        result = model.transcribe(audio_path)

        os.remove(audio_path)
        return result["text"]

    except Exception as e:
        logging.error(f"❌ Error during audio extraction/transcription: {e}")
        return ""


# --- Main: Analyze video ---
def analyze_video(file):
    if not file or file.filename == "":
        return render_template("result.html", summary="❌ No video uploaded.", mode="video"), 400

    filename = secure_filename(file.filename)
    save_path = os.path.join('static', 'uploads', filename)
    file.save(save_path)

    summary = ""
    try:
        # Extract frames
        frame_paths = extract_frames(save_path)

        # Describe frames using threads (faster than sequential)
        descriptions = [None] * len(frame_paths)
        with ThreadPoolExecutor(max_workers=5) as executor:  # adjust max_workers based on your API keys
            # future_to_frame = {executor.submit(describe_frame, f): f for f in frame_paths}
            # for future in as_completed(future_to_frame):
            #     desc = future.result()
            #     descriptions.append(desc)
            future_to_index = {executor.submit(describe_frame, f): i for i, f in enumerate(frame_paths)}
            for future in as_completed(future_to_index):
                i = future_to_index[future]
                descriptions[i] = future.result()  # Store result at correct index

        # Clean up frame images
        for path in frame_paths:
            if os.path.exists(path):
                try:
                    os.remove(path)
                except PermissionError:
                    logging.warning(f"⚠️ Could not delete frame file: {path}")

        # Extract audio & transcription
        transcript = extract_audio_and_transcribe(save_path)

        # Prepare prompt for Gemini summarization
        prompt = f"""You are analyzing a video. Visual Descriptions: {chr(10).join(descriptions)}
Audio Transcription: {transcript}. Based on these, describe the video context, type, notable people, mood, and any other insights. Provide a structured, natural, readable summary without markdown or symbols."""

        response = None
        for attempt in range(5):
            try:
                genai.configure(api_key=api_manager.current_key)
                model = genai.GenerativeModel("models/gemini-2.0-flash")
                response = model.generate_content(prompt)
                break
            except google.api_core.exceptions.ResourceExhausted:
                logging.warning("⚠️ API rate limit hit. Switching key...")
                api_manager.get_next_key()
                time.sleep(1)
            except Exception as e:
                logging.error(f"❌ Video summarization failed: {e}")
                summary = f"❌ Video summarization failed: {str(e)}"
                break
        else:
            summary = "❌ Failed to summarize video after multiple attempts due to API issues."

        if response:
            summary = response.text

    except Exception as e:
        summary = f"❌ An error occurred during video processing: {str(e)}"

    # Return result.html with video preview and summary
    return render_template(
        "result.html",
        summary=summary,
        video_url=url_for('static', filename=f'uploads/{filename}'),
        mode="video"
    )


@analysis_bp.route('/download_summary')
def download_summary():
    summary = session.get("summary", "")
    if not summary:
        return "No summary available to download."

    doc = Document()
    doc.add_heading('AI Multimedia Analyzer Summary', level=1)
    doc.add_paragraph(summary)

    upload_folder = os.path.join(analysis_bp.root_path, '..', 'static', 'uploads')
    os.makedirs(upload_folder, exist_ok=True)
    doc_path = os.path.join(upload_folder, 'summary.docx')
    doc.save(doc_path)

    @after_this_request
    def remove_file(response):
        try:
            os.remove(doc_path)
        except Exception as e:
            print("Error deleting file:", e)
        return response

    return send_file(doc_path, as_attachment=True)